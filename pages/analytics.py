import streamlit as st
import pandas as pd
from docx import Document
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from sklearn.linear_model import LinearRegression
from docx.shared import Inches
from sklearn.ensemble import RandomForestClassifier
from sklearn.cluster import KMeans
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
from sklearn.decomposition import PCA
from scipy.stats import zscore
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from wordcloud import WordCloud
from textblob import TextBlob
import os
from scipy.stats import ttest_ind, chi2_contingency 
from sklearn.cluster import AgglomerativeClustering
from sklearn.preprocessing import StandardScaler
from scipy import stats



st.set_page_config(page_title="Analytics", page_icon="🔍",initial_sidebar_state="collapsed")
with open('header.html', 'r') as file:
        html_content = file.read()
        st.markdown(html_content, unsafe_allow_html=True)


# Create a Word document for the report
doc = Document()
doc.add_heading("Data Analytics Report", 0)

# Function for anomaly detection
def perform_anomaly_detection(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Anomaly Detection")

    doc.add_paragraph("Performing Anomaly Detection...")

    selected_columns = st.multiselect(
        "Select numerical columns for Anomaly Detection",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 0:
        st.write("Report generated successfully!")

        doc.add_heading("Anomaly Detection", level=2)

        for col in selected_columns:
            doc.add_heading(f"Anomaly Detection for Column: {col}", level=3)

            if col in data.columns:
                column_data = data[col]

                if column_data.isnull().values.any():
                    doc.add_paragraph(
                        f"Column '{col}' contains missing values. Impute or clean the data before analysis."
                    )
                else:
                    z_scores = zscore(column_data)
                    threshold = 1
                    anomalies = column_data[abs(z_scores) > threshold]

                    plt.figure(figsize=(4, 3))
                    sns.histplot(column_data, kde=True, color='blue', label='Data Distribution')
                    sns.scatterplot(x=anomalies.index, y=anomalies.values, color='red', label='Anomalies')
                    plt.xlabel(col)
                    plt.ylabel('Frequency')
                    plt.title(f"Anomaly Detection for Column: {col}")
                    plt.legend()

                    anomaly_plot = BytesIO()
                    plt.savefig(anomaly_plot, format='png')
                    plt.close()

                    doc.add_picture(anomaly_plot)

                    doc.add_heading("Anomaly Detection Results", level=3)
                    doc.add_paragraph(f"Number of anomalies detected in column '{col}': {len(anomalies)}")
                    doc.add_paragraph(f"Anomaly detection threshold (z-score): {threshold}")
                    doc.add_paragraph("Insert additional insights or results here.")

                    doc.add_page_break()

def perform_clustering_and_classification(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Clustering and Classification")

    doc.add_paragraph("Performing Clustering and Classification...")

    selected_columns = st.selectbox(
        "Select numerical columns for Clustering and Classification",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )
    selected_column = st.selectbox(
        "Select Target columns for Clustering and Classification",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )
    if len(selected_columns) > 1:
        st.write("Report generated successfully!")

        doc.add_heading("Clustering and Classification", level=2)

        for col in selected_columns:
            doc.add_heading(f"Clustering and Classification for Column: {col}", level=3)

            if col in data.columns:
                X = data[selected_columns]
                y = data[selected_column]
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=42)

                kmeans = KMeans(n_clusters=2, random_state=42)
                kmeans.fit(X_train)
                cluster_labels = kmeans.predict(X_train)

                plt.figure(figsize=(4, 3))
                plt.scatter(X_train.iloc[:, 0], X_train.iloc[:, 1], c=cluster_labels, cmap='viridis')
                plt.xlabel(selected_columns[0])
                plt.ylabel(selected_columns[1])
                plt.title(f"K-Means Clustering for Column: {col}")

                clustering_plot = BytesIO()
                plt.savefig(clustering_plot, format='png')
                plt.close()

                doc.add_picture(clustering_plot)

                clf = RandomForestClassifier(random_state=42)
                clf.fit(X_train, y_train)
                y_pred = clf.predict(X_test)

                accuracy = accuracy_score(y_test, y_pred)
                doc.add_paragraph(f"Random Forest Classification Accuracy: {accuracy:.2f}")

                doc.add_page_break()




# Function for correlation analysis
def perform_correlation_analysis(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Correlation Analysis")

    doc.add_paragraph("Performing Correlation Analysis...")

    selected_columns = st.multiselect(
        "Select numerical columns for Correlation Analysis",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 1:
        st.write("Report generated successfully!")

        doc.add_heading("Correlation Analysis", level=2)

        for col in selected_columns:
            doc.add_heading(f"Correlation Analysis for Column: {col}", level=3)

            if col in data.columns:
                correlation_matrix = data[selected_columns].corr()

                plt.figure(figsize=(4, 3))
                sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', linewidths=0.5)
                plt.title(f"Correlation Heatmap for Column: {col}")

                heatmap_image = BytesIO()
                plt.savefig(heatmap_image, format='png')
                plt.close()

                doc.add_picture(heatmap_image)

                doc.add_page_break()

# Function for descriptive statistics
def perform_descriptive_statistics(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Descriptive Statistics")

    doc.add_paragraph("Performing Descriptive Statistics...")

    selected_columns = st.multiselect(
        "Select numerical columns for Descriptive Statistics",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 0:
        st.write("Report generated successfully!")

        doc.add_heading("Descriptive Statistics", level=2)

        for col in selected_columns:
            doc.add_heading(f"Descriptive Statistics for Column: {col}", level=3)

            if col in data.columns:
                desc_stats = data[col].describe()
                doc.add_paragraph(desc_stats.to_string())

                doc.add_page_break()

# Function for hypothesis testing
def perform_hypothesis_testing(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Hypothesis Testing")

    doc.add_paragraph("Performing Hypothesis Testing...")

    group_column = st.selectbox(
        "Select a categorical column for Hypothesis Testing",
        data.select_dtypes(include=[object]).columns,  # Replace np.object with object
    )
    numeric_column = st.selectbox(
        "Select a numerical column for Hypothesis Testing",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if group_column and numeric_column:
        st.write("Report generated successfully!")

        doc.add_heading("Hypothesis Testing", level=2)

        doc.add_heading(f"Hypothesis Testing for {numeric_column} by {group_column}", level=3)

        if group_column in data.columns and numeric_column in data.columns:
            group_values = data[group_column].unique()
            group_data = []

            for group in group_values:
                group_data.append(data[data[group_column] == group][numeric_column])

            if len(group_values) == 2:
                t_stat, p_value = ttest_ind(group_data[0], group_data[1])
                doc.add_paragraph(f"Two-sample t-test for {numeric_column} by {group_column}:")
                doc.add_paragraph(f"T-statistic: {t_stat:.2f}")
                doc.add_paragraph(f"P-value: {p_value:.4f}")
            else:
                contingency_table = pd.crosstab(data[group_column], data[numeric_column])
                chi2, p, _, _ = chi2_contingency(contingency_table)
                doc.add_paragraph(f"Chi-squared test for independence between {numeric_column} and {group_column}:")
                doc.add_paragraph(f"Chi-squared statistic: {chi2:.2f}")
                doc.add_paragraph(f"P-value: {p:.4f}")

            doc.add_page_break()


# Function for text analytics
# Function for text analytics
def perform_text_analytics(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Text Analytics")

    doc.add_paragraph("Performing Text Analytics...")

    text_column = st.selectbox(
        "Select a text column for Text Analytics",
        data.select_dtypes(include=[object]).columns,  # Change np.object to object
    )

    if text_column:
        st.write("Report generated successfully!")

        doc.add_heading("Text Analytics", level=2)

        doc.add_heading(f"Text Analytics for Column: {text_column}", level=3)

        if text_column in data.columns:
            text_data = data[text_column]
            text_data = text_data.dropna()

            # Perform text analysis on the selected column
            word_cloud = WordCloud(width=800, height=400, max_words=200, background_color="white").generate(
                " ".join(text_data)
            )

            plt.figure(figsize=(4, 3))
            plt.imshow(word_cloud, interpolation="bilinear")
            plt.axis("off")
            plt.title(f"Word Cloud for Column: {text_column}")

            word_cloud_image = BytesIO()
            plt.savefig(word_cloud_image, format='png')
            plt.close()

            doc.add_picture(word_cloud_image)

            doc.add_heading("Sentiment Analysis", level=3)

            sentiments = []
            for text in text_data:
                blob = TextBlob(text)
                sentiment_score = blob.sentiment.polarity
                sentiments.append(sentiment_score)

            sentiment_df = pd.DataFrame(sentiments, columns=["Sentiment Score"])

            doc.add_paragraph(sentiment_df.describe().to_string())

            doc.add_page_break()

# Function for Principal Component Analysis
def pca(data, doc):
    st.title("Data Analytics Menu")
    st.subheader("Principal Component Analysis")

    doc.add_paragraph("Performing PCA...")

    text_column = st.multiselect(
        "Select a text column for PCA",
        data.columns,  # Change np.object to object
    )

    if text_column:
        st.write("Report generated successfully!")

        doc.add_heading("Principal Component Analysis", level=2)

        doc.add_heading(f"PCA for Column: {text_column}", level=3)

        if len(text_column) > 1:
            st.write("Report generated successfully!")

            doc.add_heading("Principal Component Analysis (PCA)", level=2)

            for col in text_column:
                doc.add_heading(f"PCA Analysis for Column: {col}", level=3)

                if col in data.columns:
                    X = data[text_column]
                    X = StandardScaler().fit_transform(X)

                    pca = PCA()
                    principal_components = pca.fit_transform(X)

                    explained_variance_ratio = pca.explained_variance_ratio_
                    plt.figure(figsize=(4, 3))
                    plt.bar(range(1, len(explained_variance_ratio) + 1), explained_variance_ratio, alpha=0.5, align='center')
                    plt.xlabel('Principal Component')
                    plt.ylabel('Explained Variance Ratio')
                    plt.title(f"Explained Variance Ratio for Column: {col}")

                    explained_variance_plot = BytesIO()
                    plt.savefig(explained_variance_plot, format='png')
                    plt.close()

                    doc.add_picture(explained_variance_plot)
                    doc.add_paragraph("Principal Components:")
                    doc.add_paragraph(str(pca.components_))

                    doc.add_page_break()


def regression_analysis(data,doc):
    st.subheader("Regression Analysis")

    doc.add_paragraph("Performing Regression Analysis...")

    selected_columns = st.multiselect(
        "Select numerical columns for Regression Analysis",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 0:
        st.write("Report generated successfully!")

        for col in selected_columns:
            doc.add_heading(f"Regression Analysis for Column: {col}", level=2)

            target_variable = st.selectbox(f"Select the target variable for {col}", selected_columns)

            if target_variable != col:
                doc.add_paragraph(
                    f"Performing Linear Regression for {col} as the predictor and {target_variable} as the target variable..."
                )

                X = data[[col]]
                y = data[target_variable]
                X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

                model = LinearRegression()
                model.fit(X_train, y_train)

                r_squared = model.score(X_test, y_test)

                doc.add_paragraph(f"R-squared: {r_squared:.2f}")

                doc.add_page_break()

def time_series_analysis(data,doc):
    st.subheader("Time Series Analysis")

    doc.add_paragraph("Performing Time Series Analysis...")

    numerical_datetime_columns = [col for col in data.columns if (
        data[col].dtype in [np.float64, np.int64] or 'datetime' in str(data[col].dtype).lower()
    )]

    selected_columns = st.multiselect(
        "Select numerical and datetime columns for Time Series Analysis",
        numerical_datetime_columns,
    )

    if len(selected_columns) > 1:
        st.write("Report generated successfully!")

        doc.add_heading("Time Series Analysis", level=2)

        for col in selected_columns:
            doc.add_heading(f"Time Series Analysis for Column: {col}", level=3)

            if col in data.columns:
                time_series_data = data[[col]]

                if time_series_data.isnull().values.any():
                    doc.add_paragraph(
                        f"Column '{col}' contains missing values. Impute or clean the data before analysis."
                    )
                else:
                    plt.figure(figsize=(4, 3))
                    plt.plot(time_series_data)
                    plt.xlabel('Time')
                    plt.ylabel(col)
                    plt.title(f"Time Series Plot for Column: {col}")

                    time_series_plot = BytesIO()
                    plt.savefig(time_series_plot, format='png')
                    plt.close()

                    doc.add_picture(time_series_plot)

                    doc.add_heading("ARIMA Modeling and Forecasting", level=3)
                    doc.add_paragraph(
                        "ARIMA (AutoRegressive Integrated Moving Average) modeling is commonly used for time series analysis."
                    )

                    doc.add_paragraph("ARIMA Modeling Results")

                    doc.add_page_break()

# Function for Customer Segmentation and Analysis using Agglomerative Clustering
def perform_customer_segmentation(data, doc):
    st.title("Customer Segmentation and Analysis")
    st.subheader("Understanding Your Customer Base")

    doc.add_paragraph("Performing Customer Segmentation and Analysis...")

    selected_columns = st.multiselect(
        "Select numerical columns for Customer Segmentation",
        data.select_dtypes(include=[np.int64, np.float64]).columns,
    )

    if len(selected_columns) > 1:
        st.write("Report generated successfully!")

        doc.add_heading("Customer Segmentation", level=2)

        for col in selected_columns:
            doc.add_heading(f"Customer Segmentation for Column: {col}", level=3)

            if col in data.columns:
                # Standardize the data
                X = data[selected_columns]
                X_std = StandardScaler().fit_transform(X)

                # Apply Agglomerative Hierarchical Clustering
                n_clusters = 3  # You can adjust the number of clusters as needed
                clustering = AgglomerativeClustering(n_clusters=n_clusters)
                cluster_labels = clustering.fit_predict(X_std)

                # Visualize the clusters using the first two dimensions (you can choose different dimensions)
                plt.figure(figsize=(4, 3))
                plt.scatter(X_std[:, 0], X_std[:, 1], c=cluster_labels, cmap='viridis')
                plt.xlabel('Dimension 1')
                plt.ylabel('Dimension 2')
                plt.title(f"Customer Segmentation for {col}")
                clustering_plot = BytesIO()
                plt.savefig(clustering_plot, format='png')
                plt.close()

                doc.add_picture(clustering_plot)

                doc.add_heading("Cluster Analysis Results", level=3)
                doc.add_paragraph(f"Number of Clusters: {n_clusters}")
                doc.add_page_break()


# Main App
def main():
    st.markdown("<div style='text-align:center;'><h1>Data Analytics</h1></div>", unsafe_allow_html=True)
    st.markdown("<div style='text-align:center;'>Upload CSV File</div>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader("", type=["csv"])

    if uploaded_file is not None:
        data = pd.read_csv(uploaded_file)

        st.write("Data Preview:")
        st.write(data.head())

        doc.add_paragraph("Data Preview:")
        doc.add_paragraph(data.head().to_string())

        st.write("Data Info:")
        st.write(f"Number of Rows: {data.shape[0]}")
        st.write(f"Number of Columns: {data.shape[1]}")

        # Save the uploaded data as a CSV file for reference in the report
        data.to_csv("uploaded_data.csv", index=False)

        st.success("Data uploaded successfully!")

        st.title("Data Analytics Menu")
        st.subheader("Choose an Analytics Task")

        analytics_task = st.selectbox(
                "Select an Analytics Task",
                ["Anomaly Detection", "Customer Segmentation", "Correlation Analysis",
                "Descriptive Statistics", "Hypothesis Testing", "Text Analytics","PCA","Regression Analysis","Time Series Analysis"],
        )

        if analytics_task == "Anomaly Detection":
                perform_anomaly_detection(data, doc)

        elif analytics_task == "Clustering and Classification":
                perform_clustering_and_classification(data, doc)

        elif analytics_task == "Correlation Analysis":
                perform_correlation_analysis(data, doc)

        elif analytics_task == "Descriptive Statistics":
                perform_descriptive_statistics(data, doc)

        elif analytics_task == "Hypothesis Testing":
                perform_hypothesis_testing(data, doc)

        elif analytics_task == "Text Analytics":
                perform_text_analytics(data, doc)

        elif analytics_task == "PCA":
                pca(data, doc)

        elif analytics_task == "Regression Analysis":
                regression_analysis(data, doc)

        elif analytics_task == "Time Series Analysis":
                time_series_analysis(data, doc)

        elif analytics_task == "Customer Segmentation":
                perform_customer_segmentation(data,doc)

        else:
                st.warning("Please select at least one analysis technique.")

        directory = os.path.expanduser("~/Desktop")
        # Create the directory if it doesn't exist
        report_filename = "data_analytics_report.docx"
        report_filepath = os.path.join(directory, report_filename)

        doc.save(report_filepath)

        # Provide a link to download the report
        st.success(f"Report generated successfully! You can download the report from the following link: [Download Report]({report_filepath})")
if __name__ == "__main__":
    main()
