import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
import plotly.express as px
import os
import plotly.graph_objects as go
from wordcloud import WordCloud
import networkx as nx
import docx
from docx.shared import Inches


st.set_page_config(page_title="Visualization", page_icon="👀", initial_sidebar_state="collapsed")
with open('header.html', 'r') as file:
        header_html_content = file.read()
st.markdown(header_html_content, unsafe_allow_html=True)

st.markdown("<div style='text-align:center;'><h1>Data Analytics</h1></div>", unsafe_allow_html=True)
st.markdown("<div style='text-align:center;'>Upload CSV or XLSX File</div>", unsafe_allow_html=True)

uploaded_file = st.file_uploader("", type=["csv", "xlsx"])

if uploaded_file is not None:
    df = pd.read_csv(uploaded_file) 

    available_plots = [
        "Line Plot",
        "Scatter Plot",
        "Bar Plot (Vertical and Horizontal)",
        "Stacked Bar Plot",
        "Grouped Bar Plot",
        "Histogram",
        "Frequency Polygon",
        "Box Plot",
        "Violin Plot",
        "Density Plot",
        "Area Plot",
        "Stacked Area Plot",
        "Step Plot",
        "Pie Chart",
        "Donut Chart",
        "Polar Plot",
        "3D Line Plot",
        "3D Scatter Plot",
        "3D Surface Plot",
        "Wireframe Plot",
        "Contour Plot",
        "Filled Contour Plot",
        "Hexbin Plot (2D Histogram)",
        "Quiver Plot",
        "Streamplot",
        "Stem Plot",
        "Error Bars (Vertical and Horizontal)",
        "Candlestick Chart",
        "Radial Bar Plot",
        "Radar Chart",
        "Ternary Plot",
        "Sankey Diagram",
        "Chord Diagram",
        "Word Cloud",
        "Network Graph"
    ]

    image_dir = "images"
    os.makedirs(image_dir, exist_ok=True)


    selected_plots = st.multiselect("Select plots to visualize", available_plots)

    image_paths = []

    for plot_type in selected_plots:
        if "Line Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Line Plot", df.columns ,key="line_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Line Plot", df.columns,key="line_y_column_selectbox")
            plt.figure()
            plt.plot(df[x_column], df[y_column])
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Line Plot")
            line_plot_image_path = os.path.join(image_dir, "line_plot.png")
            plt.savefig(line_plot_image_path)
            image_paths.append(line_plot_image_path)

        elif "Scatter Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Scatter Plot", df.columns,key="scatter_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Scatter Plot", df.columns,key="scatter_y_column_selectbox")
            plt.figure()
            plt.scatter(df[x_column], df[y_column])
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Scatter Plot")
            scatter_plot_image_path = os.path.join(image_dir, "scatter_plot.png")
            plt.savefig(scatter_plot_image_path)
            image_paths.append(scatter_plot_image_path)

        elif "Bar Plot (Vertical and Horizontal)" in plot_type:
            x_column = st.selectbox("Select X-axis column for Bar Plot", df.columns,key="bar_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Bar Plot", df.columns,key="bar_y_column_selectbox")
            orientation = st.radio("Select orientation", ["vertical", "horizontal"])
            plt.figure()
            if orientation == "vertical":
                plt.bar(df[x_column], df[y_column])
                plt.xlabel(x_column)
                plt.ylabel(y_column)
            else:
                plt.barh(df[x_column], df[y_column])
                plt.xlabel(y_column)
                plt.ylabel(x_column)
            plt.title("Bar Plot")
            bar_plot_image_path = os.path.join(image_dir, "bar_plot.png")
            plt.savefig(bar_plot_image_path)
            image_paths.append(bar_plot_image_path)



        elif "Stacked Bar Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Stacked Bar Plot", df.columns,key="sbar_x_column_selectbox")
            y_columns = st.multiselect("Select Y-axis columns for Stacked Bar Plot", df.columns,key="sbar_y_column_selectbox")
            stacked_bar_df = df[y_columns].copy()
            stacked_bar_df[x_column] = df[x_column]
            stacked_bar_df.set_index(x_column, inplace=True)
            stacked_bar_df.plot(kind="bar", stacked=True)
            plt.xlabel(x_column)
            plt.ylabel("Value")
            plt.title("Stacked Bar Plot")
            stacked_bar_plot_image_path = os.path.join(image_dir, "stacked_bar_plot.png")
            plt.savefig(stacked_bar_plot_image_path)
            image_paths.append(stacked_bar_plot_image_path)

        elif "Grouped Bar Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Grouped Bar Plot", df.columns,key="gbar_x_column_selectbox")
            y_columns = st.multiselect("Select Y-axis columns for Grouped Bar Plot", df.columns,key="gbar_y_column_selectbox")
            grouped_bar_df = df[y_columns].copy()
            grouped_bar_df[x_column] = df[x_column]
            grouped_bar_df.set_index(x_column, inplace=True)
            grouped_bar_df.plot(kind="bar")
            plt.xlabel(x_column)
            plt.ylabel("Value")
            plt.title("Grouped Bar Plot")
            grouped_bar_plot_image_path = os.path.join(image_dir, "grouped_bar_plot.png")
            plt.savefig(grouped_bar_plot_image_path)
            image_paths.append(grouped_bar_plot_image_path)

    
        elif "Histogram" in plot_type:
            column = st.selectbox("Select column for Histogram", df.columns,key="histogram_x_column_selectbox")
            plt.figure()
            plt.hist(df[column], bins=20, edgecolor='k')
            plt.xlabel(column)
            plt.ylabel("Frequency")
            plt.title("Histogram")
            hist_image_path = os.path.join(image_dir, "histogram.png")
            plt.savefig(hist_image_path)
            image_paths.append(hist_image_path)



        elif "Frequency Polygon" in plot_type:
            column = st.selectbox("Select column for Frequency Polygon", df.columns,key="fpolygon_x_column_selectbox")
            plt.figure()
            plt.plot(df[column], pd.Series(df[column]).value_counts(sort=False), marker='o')
            plt.xlabel(column)
            plt.ylabel("Frequency")
            plt.title("Frequency Polygon")
            freq_poly_image_path = os.path.join(image_dir, "frequency_polygon.png")
            plt.savefig(freq_poly_image_path)
            image_paths.append(freq_poly_image_path)


        elif "Box Plot" in plot_type:
            column = st.selectbox("Select column for Box Plot", df.columns,key="box_x_column_selectbox")
            plt.figure()
            plt.boxplot(df[column], vert=False)
            plt.xlabel(column)
            plt.title("Box Plot")
            box_plot_image_path = os.path.join(image_dir, "box_plot.png")
            plt.savefig(box_plot_image_path)
            image_paths.append(box_plot_image_path)

        elif "Violin Plot" in plot_type:
            column = st.selectbox("Select column for Violin Plot", df.columns,key="violin_x_column_selectbox")
            plt.figure()
            sns.violinplot(data=df, y=column)
            plt.ylabel(column)
            plt.title("Violin Plot")
            violin_plot_image_path = os.path.join(image_dir, "violin_plot.png")
            plt.savefig(violin_plot_image_path)
            image_paths.append(violin_plot_image_path)

        elif "Density Plot" in plot_type:
            column = st.selectbox("Select column for Density Plot", df.columns,key="density_x_column_selectbox")
            plt.figure()
            sns.kdeplot(df[column], shade=True)
            plt.xlabel(column)
            plt.ylabel("Density")
            plt.title("Density Plot")
            density_plot_image_path = os.path.join(image_dir, "density_plot.png")
            plt.savefig(density_plot_image_path)
            image_paths.append(density_plot_image_path)


        elif "Area Plot" in plot_type:
            column = st.selectbox("Select column for Area Plot", df.columns,key="areaa_x_colunnmn_selectbox")
            plt.figure()
            plt.fill_between(df.index, df[column])
            plt.xlabel("Index")
            plt.ylabel(column)
            plt.title("Area Plot")
            area_plot_image_path = os.path.join(image_dir, "area_plot.png")
            plt.savefig(area_plot_image_path)
            image_paths.append(area_plot_image_path)

        elif "Stacked Area Plot" in plot_type:
            column = st.selectbox("Select column for Stacked Area Plot", df.columns,key="sarea_x_column_selectbox")
            stacked_area_df = df.copy()
            stacked_area_df.set_index(column, inplace=True)
            stacked_area_df.plot(kind="area", stacked=True)
            plt.xlabel(column)
            plt.ylabel("Value")
            plt.title("Stacked Area Plot")
            stacked_area_plot_image_path = os.path.join(image_dir, "stacked_area_plot.png")
            plt.savefig(stacked_area_plot_image_path)
            image_paths.append(stacked_area_plot_image_path)

        elif "Step Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Step Plot", df.columns,key="step_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Step Plot", df.columns,key="step_y_column_selectbox")
            plt.figure()
            plt.step(df[x_column], df[y_column])
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Step Plot")
            step_plot_image_path = os.path.join(image_dir, "step_plot.png")
            plt.savefig(step_plot_image_path)
            image_paths.append(step_plot_image_path)

        elif "Pie Chart" in plot_type:
            column = st.selectbox("Select column for Pie Chart", df.columns,key="pie_x_column_selectbox")
            plt.figure()
            plt.pie(df[column], labels=df.index, autopct="%1.1f%%", startangle=140)
            plt.title("Pie Chart")
            pie_chart_image_path = os.path.join(image_dir, "pie_chart.png")
            plt.savefig(pie_chart_image_path)
            image_paths.append(pie_chart_image_path)


        elif "Donut Chart" in plot_type:
            column = st.selectbox("Select column for Donut Chart", df.columns,key="donut_x_column_selectbox")
            plt.figure()
            plt.pie(df[column], labels=df.index, autopct="%1.1f%%", startangle=140, wedgeprops={"edgecolor": "white"})
            centre_circle = plt.Circle((0, 0), 0.70, fc='white')
            fig = plt.gcf()
            fig.gca().add_artist(centre_circle)
            plt.title("Donut Chart")
            donut_chart_image_path = os.path.join(image_dir, "donut_chart.png")
            plt.savefig(donut_chart_image_path)
            image_paths.append(donut_chart_image_path)

        elif "Polar Plot" in plot_type:
            column = st.selectbox("Select column for Polar Plot", df.columns,key="polar_x_column_selectbox")
            plt.figure()
            plt.polar(df.index, df[column])
            plt.title("Polar Plot")
            polar_plot_image_path = os.path.join(image_dir, "polar_plot.png")
            plt.savefig(polar_plot_image_path)
            image_paths.append(polar_plot_image_path)


        elif "3D Line Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for 3D Line Plot", df.columns,key="3d_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for 3D Line Plot", df.columns,key="3d_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for 3D Line Plot", df.columns,key="3d_z_column_selectbox")
            fig = plt.figure()
            ax = fig.add_subplot(111, projection="3d")
            ax.plot(df[x_column], df[y_column], df[z_column])
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_zlabel(z_column)
            plt.title("3D Line Plot")
            line3d_plot_image_path = os.path.join(image_dir, "3d_line_plot.png")
            plt.savefig(line3d_plot_image_path)
            image_paths.append(line3d_plot_image_path)


        elif "3D Scatter Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for 3D Scatter Plot", df.columns,key="3ds_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for 3D Scatter Plot", df.columns,key="3ds_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for 3D Scatter Plot", df.columns,key="3ds_z_column_selectbox")
            fig = plt.figure()
            ax = fig.add_subplot(111, projection="3d")
            ax.scatter(df[x_column], df[y_column], df[z_column])
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_zlabel(z_column)
            plt.title("3D Scatter Plot")
            scatter3d_plot_image_path = os.path.join(image_dir, "3d_scatter_plot.png")
            plt.savefig(scatter3d_plot_image_path)
            image_paths.append(scatter3d_plot_image_path)


        elif "3D Surface Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for 3D Surface Plot", df.columns,key="3dc_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for 3D Surface Plot", df.columns,key="3dc_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for 3D Surface Plot", df.columns,key="3dc_z_column_selectbox")
            fig = plt.figure()
            ax = fig.add_subplot(111, projection="3d")
            ax.plot_surface(df[x_column], df[y_column], df[z_column], cmap="viridis")
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_zlabel(z_column)
            plt.title("3D Surface Plot")
            surface3d_plot_image_path = os.path.join(image_dir, "3d_surface_plot.png")
            plt.savefig(surface3d_plot_image_path)
            image_paths.append(surface3d_plot_image_path)


        elif "Wireframe Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Wireframe Plot", df.columns,key="wire_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Wireframe Plot", df.columns,key="wire_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for Wireframe Plot", df.columns,key="wire_z_column_selectbox")
            fig = plt.figure()
            ax = fig.add_subplot(111, projection="3d")
            ax.plot_wireframe(df[x_column], df[y_column], df[z_column])
            ax.set_xlabel(x_column)
            ax.set_ylabel(y_column)
            ax.set_zlabel(z_column)
            plt.title("Wireframe Plot")
            wireframe_plot_image_path = os.path.join(image_dir, "wireframe_plot.png")
            plt.savefig(wireframe_plot_image_path)
            image_paths.append(wireframe_plot_image_path)


        elif "Contour Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Contour Plot", df.columns,key="con_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Contour Plot", df.columns,key="con_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for Contour Plot", df.columns,key="con_z_column_selectbox")
            plt.figure()
            plt.contour(df[x_column], df[y_column], df[z_column], cmap="viridis")
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Contour Plot")
            contour_plot_image_path = os.path.join(image_dir, "contour_plot.png")
            plt.savefig(contour_plot_image_path)
            image_paths.append(contour_plot_image_path)


        elif "Filled Contour Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Filled Contour Plot", df.columns,key="fcon_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Filled Contour Plot", df.columns,key="fcon_y_column_selectbox")
            z_column = st.selectbox("Select Z-axis column for Filled Contour Plot", df.columns,key="fcon_z_column_selectbox")
            plt.figure()
            plt.contourf(df[x_column], df[y_column], df[z_column], cmap="viridis")
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Filled Contour Plot")
            filled_contour_plot_image_path = os.path.join(image_dir, "filled_contour_plot.png")
            plt.savefig(filled_contour_plot_image_path)
            image_paths.append(filled_contour_plot_image_path)


        elif "Hexbin Plot (2D Histogram)" in plot_type:
            x_column = st.selectbox("Select X-axis column for Hexbin Plot", df.columns,key="hex_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Hexbin Plot", df.columns,key="hex_y_column_selectbox")
            plt.figure()
            plt.hexbin(df[x_column], df[y_column], gridsize=20, cmap="viridis")
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Hexbin Plot")
            hexbin_plot_image_path = os.path.join(image_dir, "hexbin_plot.png")
            plt.savefig(hexbin_plot_image_path)
            image_paths.append(hexbin_plot_image_path)


        elif "Quiver Plot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Quiver Plot", df.columns,key="q_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Quiver Plot", df.columns,key="q_y_column_selectbox")
            u_column = st.selectbox("Select U-axis column for Quiver Plot", df.columns,key="q_z_column_selectbox")
            v_column = st.selectbox("Select V-axis column for Quiver Plot", df.columns,key="q_r_column_selectbox")
            plt.figure()
            plt.quiver(df[x_column], df[y_column], df[u_column], df[v_column])
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Quiver Plot")
            quiver_plot_image_path = os.path.join(image_dir, "quiver_plot.png")
            plt.savefig(quiver_plot_image_path)
            image_paths.append(quiver_plot_image_path)


        elif "Streamplot" in plot_type:
            x_column = st.selectbox("Select X-axis column for Streamplot", df.columns,key="st_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Streamplot", df.columns,key="st_y_column_selectbox")
            u_column = st.selectbox("Select U-axis column for Streamplot", df.columns,key="st_z_column_selectbox")
            v_column = st.selectbox("Select V-axis column for Streamplot", df.columns,key="st_r_column_selectbox")
            plt.figure()
            plt.streamplot(df[x_column], df[y_column], df[u_column], df[v_column])
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Streamplot")
            streamplot_image_path = os.path.join(image_dir, "streamplot.png")
            plt.savefig(streamplot_image_path)
            image_paths.append(streamplot_image_path)


        elif "Stem Plot" in plot_type:
            column = st.selectbox("Select column for Stem Plot", df.columns,key="stem_x_column_selectbox")
            plt.figure()
            plt.stem(df[column])
            plt.xlabel("Index")
            plt.ylabel(column)
            plt.title("Stem Plot")
            stem_plot_image_path = os.path.join(image_dir, "stem_plot.png")
            plt.savefig(stem_plot_image_path)
            image_paths.append(stem_plot_image_path)


        elif "Error Bars (Vertical and Horizontal)" in plot_type:
            x_column = st.selectbox("Select X-axis column for Error Bars Plot", df.columns,key="er_x_column_selectbox")
            y_column = st.selectbox("Select Y-axis column for Error Bars Plot", df.columns,key="er_y_column_selectbox")
            error_column = st.selectbox("Select Error column for Error Bars Plot", df.columns,key="er_z_column_selectbox")
            plt.figure()
            plt.errorbar(df[x_column], df[y_column], yerr=df[error_column], fmt="o")
            plt.xlabel(x_column)
            plt.ylabel(y_column)
            plt.title("Error Bars Plot")
            error_bars_plot_image_path = os.path.join(image_dir, "error_bars_plot.png")
            plt.savefig(error_bars_plot_image_path)
            image_paths.append(error_bars_plot_image_path)


        elif "Candlestick Chart" in plot_type:
            candlestick_df = df.copy()
            candlestick_df.reset_index(inplace=True)
            candlestick_df["Date"] = pd.to_datetime(candlestick_df["Date"])
            fig = go.Figure(data=[go.Candlestick(x=candlestick_df["Date"],
                                                  open=candlestick_df["Open"],
                                                  high=candlestick_df["High"],
                                                  low=candlestick_df["Low"],
                                                  close=candlestick_df["Close"])])
            fig.update_layout(title="Candlestick Chart")
            candlestick_chart_image_path = os.path.join(image_dir, "candlestick_chart.png")
            fig.write_image(candlestick_chart_image_path)
            image_paths.append(candlestick_chart_image_path)


        elif "Radial Bar Plot" in plot_type:
            theta_column = st.selectbox("Select theta column for Radial Bar Plot", df.columns,key="rbar_x_column_selectbox")
            r_column = st.selectbox("Select r column for Radial Bar Plot", df.columns)
            fig = px.bar_polar(df, r=r_column, theta=theta_column, title="Radial Bar Plot")
            radial_bar_plot_image_path = os.path.join(image_dir, "radial_bar_plot.png")
            fig.write_image(radial_bar_plot_image_path)
            image_paths.append(radial_bar_plot_image_path)

        elif "Radar Chart" in plot_type:
            column_names = st.multiselect("Select columns for Radar Chart", df.columns,key="radar_x_column_selectbox")
            fig = px.line_polar(df, r=column_names, theta=df.index, title="Radar Chart")
            radar_chart_image_path = os.path.join(image_dir, "radar_chart.png")
            fig.write_image(radar_chart_image_path)
            image_paths.append(radar_chart_image_path)

        elif "Ternary Plot" in plot_type:
            a_column = st.selectbox("Select A-axis column for Ternary Plot", df.columns,key="ter_x_column_selectbox")
            b_column = st.selectbox("Select B-axis column for Ternary Plot", df.columns,key="ter_y_column_selectbox")
            c_column = st.selectbox("Select C-axis column for Ternary Plot", df.columns,key="ter_z_column_selectbox")
            fig = px.scatter_ternary(df, a=a_column, b=b_column, c=c_column, title="Ternary Plot")
            ternary_plot_image_path = os.path.join(image_dir, "ternary_plot.png")
            fig.write_image(ternary_plot_image_path)
            image_paths.append(ternary_plot_image_path)


        elif "Sankey Diagram" in plot_type:
            source_column = st.selectbox("Select Source column for Sankey Diagram", df.columns,key="sank_x_column_selectbox")
            target_column = st.selectbox("Select Target column for Sankey Diagram", df.columns,key="sank_y_column_selectbox")
            value_column = st.selectbox("Select Value column for Sankey Diagram", df.columns,key="sank_z_column_selectbox")
            fig = go.Figure(go.Sankey(
                node=dict(
                    pad=15,
                    thickness=20,
                    line=dict(color="black", width=0.5),
                    label=df[source_column].tolist() + df[target_column].tolist(),
                ),
                link=dict(
                    source=df[source_column].tolist(),
                    target=df[target_column].tolist(),
                    value=df[value_column].tolist(),
                )
            ))
            fig.update_layout(title="Sankey Diagram")
            sankey_diagram_image_path = os.path.join(image_dir, "sankey_diagram.png")
            fig.write_image(sankey_diagram_image_path)
            image_paths.append(sankey_diagram_image_path)

        elif "Chord Diagram" in plot_type:
            chord_matrix = df.corr()
            fig = go.Figure(go.Chord(
                matrix=chord_matrix.values,
                labels=df.columns.tolist(),
            ))
            fig.update_layout(title="Chord Diagram")
            chord_diagram_image_path = os.path.join(image_dir, "chord_diagram.png")
            fig.write_image(chord_diagram_image_path)
            image_paths.append(chord_diagram_image_path)

        elif "Word Cloud" in plot_type:
            text_column = st.selectbox("Select Text column for Word Cloud", df.columns,key="word_x_column_selectbox")
            wordcloud = WordCloud(width=800, height=400, background_color="white").generate(" ".join(df[text_column]))
            plt.figure(figsize=(10, 5))
            plt.imshow(wordcloud, interpolation="bilinear")
            plt.axis("off")
            plt.title("Word Cloud")
            wordcloud_image_path = os.path.join(image_dir, "wordcloud.png")
            plt.savefig(wordcloud_image_path)
            image_paths.append(wordcloud_image_path)

        elif "Network Graph" in plot_type:
            source_column = st.selectbox("Select Source column for Network Graph", df.columns,key="netw_x_column_selectbox")
            target_column = st.selectbox("Select Target column for Network Graph", df.columns,key="netw_y_column_selectbox")
            G = nx.from_pandas_edgelist(df, source=source_column, target=target_column)
            pos = nx.spring_layout(G)
            edge_x = []
            edge_y = []
            for edge in G.edges():
                x0, y0 = pos[edge[0]]
                x1, y1 = pos[edge[1]]
                edge_x.append(x0)
                edge_x.append(x1)
                edge_x.append(None)
                edge_y.append(y0)
                edge_y.append(y1)
                edge_y.append(None)
            edge_trace = go.Scatter(
                x=edge_x, y=edge_y,
                line=dict(width=0.5, color="#888"),
                hoverinfo="none",
                mode="lines")
            node_x = []
            node_y = []
            for node in G.nodes():
                x, y = pos[node]
                node_x.append(x)
                node_y.append(y)
            node_trace = go.Scatter(
                x=node_x, y=node_y,
                mode="markers",
                hoverinfo="text",
                marker=dict(
                    showscale=True,
                    colorscale="YlGnBu",
                    size=10,
                    colorbar=dict(
                        thickness=15,
                        title="Node Connections",
                        xanchor="left",
                        titleside="right"
                    )
                )
            )

            node_adjacencies = []
            node_text = []
            for node, adjacencies in enumerate(G.adjacency()):
                node_adjacencies.append(len(adjacencies[1]))
                node_text.append(f"Connections: {len(adjacencies[1])}")

            node_trace.marker.color = node_adjacencies
            node_trace.text = node_text

            fig = go.Figure(data=[edge_trace, node_trace],
                            layout=px.Layout(
                                showlegend=False,
                                hovermode="closest",
                                margin=dict(b=0, l=0, r=0, t=0),
                            ))
            fig.update_layout(title="Network Graph")
            network_graph_image_path = os.path.join(image_dir, "network_graph.png")
            fig.write_image(network_graph_image_path)
            image_paths.append(network_graph_image_path)







if st.button("Generate and Save Graph to Word"):
    doc = docx.Document()

    for plot_type, image_path in zip(selected_plots, image_paths):
        st.image(image_path)
        # Add a title for the plot
        doc.add_heading(plot_type, level=1)
        
        # Add the plot image to the document
        doc.add_picture(image_path, width=Inches(6))
        directory = os.path.expanduser("~/Desktop")
        # Create the directory if it doesn't exist
        report_filename = "data_visualization_report.docx"
        report_filepath = os.path.join(directory, report_filename)
        doc.save(report_filepath)
        # Provide a link to download the report
    st.success(f"Report generated successfully! You can download the report from the following link: [Download Report]({report_filepath})")
    # Save the document
    st.success("Graphs have been saved to 'graphs.docx'")